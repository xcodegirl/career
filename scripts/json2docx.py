"""
Generate Word document resume from JSON source.
Usage: python json2docx.py input.json [output.docx]
"""
import os
import sys

from resume_builder_common import get_output_path, load_resume_json, build_sections

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def _add_hyperlink(paragraph, url, text):
    """Embed a real Word hyperlink into an existing paragraph."""
    r_id = paragraph.part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '3D5AF1')
    rPr.append(color_el)
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def add_section_header(doc, title):
    """Add a formatted section header with accent color."""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(6)

    run = heading.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 60, 90)  # Accent blue
    run.font.name = 'Calibri'


# =====================================================================
# SECTION BUILDERS
# =====================================================================

def add_summary_section(doc, resume_data):
    """Add professional summary."""
    if not resume_data.get("summary"):
        return
    add_section_header(doc, "Summary")
    doc.add_paragraph(resume_data["summary"])


def add_ai_expertise_section(doc, resume_data):
    """Add AI expertise section if present."""
    if not resume_data.get("ai_expertise"):
        return

    add_section_header(doc, "AI Expertise")

    ai = resume_data["ai_expertise"]

    if ai.get("daily_practice"):
        para = doc.add_paragraph()
        para.add_run("Daily practice:").bold = True
        for item in ai["daily_practice"]:
            doc.add_paragraph(item, style='List Bullet')

    if ai.get("teaching"):
        para = doc.add_paragraph()
        para.add_run("What I teach about AI:").bold = True
        for item in ai["teaching"]:
            doc.add_paragraph(item, style='List Bullet')

    if ai.get("assessment"):
        para = doc.add_paragraph()
        para.add_run("AI and assessment design (from training):").bold = True
        for item in ai["assessment"]:
            doc.add_paragraph(item, style='List Bullet')


def add_experience_section(doc, resume_data):
    """Add experience section."""
    if not resume_data.get("experience"):
        return
    add_section_header(doc, "Experience")
    for job in resume_data["experience"]:
        job_para = doc.add_paragraph(style='List Bullet')
        title_run = job_para.add_run(f"{job.get('title', '')} - {job.get('company', '')}")
        title_run.bold = True
        title_run.font.size = Pt(11)

        date_para = doc.add_paragraph(job.get("date", ""), style='List Bullet 2')
        date_para.paragraph_format.left_indent = Inches(0.5)
        if date_para.runs:
            date_run = date_para.runs[0]
            date_run.font.italic = True
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(100, 100, 100)

        if job.get("technologies"):
            tech_para = doc.add_paragraph(style='List Bullet 2')
            tech_para.paragraph_format.left_indent = Inches(0.5)
            tech_run = tech_para.add_run("Technologies: ")
            tech_run.bold = True
            tech_para.add_run(" / ".join(job["technologies"]))

        if job.get("description"):
            desc_para = doc.add_paragraph(job.get("description"), style='List Bullet 2')
            desc_para.paragraph_format.left_indent = Inches(0.5)

        if job.get("bullets"):
            for bullet in job["bullets"]:
                bullet_para = doc.add_paragraph(bullet, style='List Bullet 2')
                bullet_para.paragraph_format.left_indent = Inches(0.5)


def add_education_section(doc, resume_data):
    """Add education section."""
    if not resume_data.get("education"):
        return
    add_section_header(doc, "Education")
    for school in resume_data["education"]:
        edu_para = doc.add_paragraph(style='List Bullet')
        degree_run = edu_para.add_run(f"{school.get('degree', '')}")
        degree_run.bold = True
        degree_run.font.size = Pt(11)
        edu_para.add_run(f" in {school.get('field', '')}")

        inst_para = doc.add_paragraph(style='List Bullet 2')
        inst_para.paragraph_format.left_indent = Inches(0.5)
        inst_run = inst_para.add_run(school.get('institution', ''))
        inst_run.italic = True

        date_para = doc.add_paragraph(school.get('date', ''), style='List Bullet 2')
        date_para.paragraph_format.left_indent = Inches(0.5)
        if date_para.runs:
            date_run = date_para.runs[0]
            date_run.font.color.rgb = RGBColor(100, 100, 100)

        if school.get('thesis'):
            thesis_para = doc.add_paragraph(style='List Bullet 2')
            thesis_para.paragraph_format.left_indent = Inches(0.5)
            thesis_para.add_run("Thesis: ").bold = True
            thesis_run = thesis_para.add_run(school['thesis'])
            thesis_run.italic = True


def add_certifications_section(doc, resume_data):
    """Add certifications section."""
    if not resume_data.get("certifications"):
        return
    add_section_header(doc, "Certifications")
    for cert in resume_data["certifications"]:
        cert_para = doc.add_paragraph(style='List Bullet')
        cert_run = cert_para.add_run(cert.get('title', ''))
        cert_run.bold = True
        if cert.get('issuer'):
            cert_para.add_run(f" - {cert.get('issuer', '')}")
        if cert.get('date'):
            cert_para.add_run(f" ({cert.get('date', '')})")
        if cert.get('credential_id'):
            cred_para = doc.add_paragraph(f"ID: {cert.get('credential_id', '')}", style='List Bullet 2')
            cred_para.paragraph_format.left_indent = Inches(0.5)


def add_awards_section(doc, resume_data):
    """Add awards section."""
    if not resume_data.get("awards"):
        return
    add_section_header(doc, "Awards & Honours")
    for award in resume_data["awards"]:
        award_para = doc.add_paragraph(style='List Bullet')
        award_run = award_para.add_run(award.get('title', ''))
        award_run.bold = True
        award_para.add_run(f" - {award.get('organization', '')} ({award.get('year', '')})")


def add_skills_section(doc, resume_data):
    """Add skills section."""
    if not resume_data.get("skills"):
        return
    add_section_header(doc, "Skills")
    skills = resume_data["skills"]
    if isinstance(skills, dict):
        for category, items in skills.items():
            cat_para = doc.add_paragraph(style='List Bullet')
            cat_run = cat_para.add_run(f"{category}:")
            cat_run.bold = True
            if isinstance(items, list):
                items_str = ', '.join(items)
            else:
                items_str = str(items)
            cat_para.add_run(f" {items_str}")


def add_projects_section(doc, resume_data):
    """Add projects section."""
    if not resume_data.get("projects"):
        return
    add_section_header(doc, "Projects")
    for project in resume_data["projects"]:
        proj_para = doc.add_paragraph(style='List Bullet')
        proj_run = proj_para.add_run(project.get('title', ''))
        proj_run.bold = True
        if project.get('date'):
            proj_para.add_run(f" ({project.get('date', '')})")

        desc_para = doc.add_paragraph(project.get('description', ''), style='List Bullet 2')
        desc_para.paragraph_format.left_indent = Inches(0.5)

        if project.get('technologies'):
            tech_str = ', '.join(project['technologies'])
            tech_para = doc.add_paragraph(f"Technologies: {tech_str}", style='List Bullet 2')
            tech_para.paragraph_format.left_indent = Inches(0.5)


def add_volunteer_section(doc, resume_data):
    """Add volunteer experience section."""
    if not resume_data.get("volunteer"):
        return
    add_section_header(doc, "Volunteer Experience")
    for vol in resume_data["volunteer"]:
        vol_para = doc.add_paragraph(style='List Bullet')
        vol_run = vol_para.add_run(vol.get('title', ''))
        vol_run.bold = True
        vol_para.add_run(f" - {vol.get('organization', '')}")

        date_para = doc.add_paragraph(vol.get('date', ''), style='List Bullet 2')
        date_para.paragraph_format.left_indent = Inches(0.5)

        desc_para = doc.add_paragraph(vol.get('description', ''), style='List Bullet 2')
        desc_para.paragraph_format.left_indent = Inches(0.5)


def add_publications_section(doc, resume_data):
    """Add publications section."""
    if not resume_data.get("publications"):
        return
    add_section_header(doc, "Publications")
    for pub in resume_data["publications"]:
        pub_para = doc.add_paragraph(style='List Bullet')
        pub_run = pub_para.add_run(pub.get('title', ''))
        pub_run.bold = True
        if pub.get('publication'):
            pub_para.add_run(f" - {pub.get('publication', '')}")
        if pub.get('date'):
            pub_para.add_run(f" ({pub.get('date', '')})")


def add_languages_section(doc, resume_data):
    """Add languages section."""
    if not resume_data.get("languages"):
        return
    add_section_header(doc, "Languages")
    for lang in resume_data["languages"]:
        lang_para = doc.add_paragraph(style='List Bullet')
        lang_para.add_run(f"{lang.get('language', '')}: {lang.get('proficiency', '')}")


def add_memberships_section(doc, resume_data):
    """Add professional memberships section."""
    if not resume_data.get("memberships"):
        return
    add_section_header(doc, "Professional Memberships")
    for membership in resume_data["memberships"]:
        mem_para = doc.add_paragraph(style='List Bullet')
        mem_run = mem_para.add_run(membership.get('organization', ''))
        mem_run.bold = True
        if membership.get('title'):
            mem_para.add_run(f" - {membership.get('title', '')}")
        if membership.get('date'):
            mem_para.add_run(f" ({membership.get('date', '')})")


def add_portfolio_section(doc, resume_data):
    """Add portfolio section."""
    if not resume_data.get("portfolio"):
        return
    add_section_header(doc, "Portfolio")
    for item in resume_data["portfolio"]:
        port_para = doc.add_paragraph(style='List Bullet')
        _add_hyperlink(port_para, item.get('url', ''), item.get('title', ''))


def add_published_games_section(doc, resume_data):
    """Add published games section."""
    if not resume_data.get("published_games"):
        return
    add_section_header(doc, "Published Games")
    for game in resume_data["published_games"]:
        platforms = ', '.join(game["platforms"]) if isinstance(game["platforms"], list) else game["platforms"]
        year = game["year"] if isinstance(game["year"], str) else ', '.join(game["year"])
        game_para = doc.add_paragraph(style='List Bullet')
        game_run = game_para.add_run(game.get('title', ''))
        game_run.bold = True
        game_para.add_run(f" ({platforms}, {year})")


# =====================================================================
# SECTION MAPPING AND ORDERING
# =====================================================================

SECTION_BUILDERS = {
    'summary': add_summary_section,
    'ai_expertise': add_ai_expertise_section,
    'experience': add_experience_section,
    'education': add_education_section,
    'certifications': add_certifications_section,
    'awards': add_awards_section,
    'skills': add_skills_section,
    'projects': add_projects_section,
    'volunteer': add_volunteer_section,
    'publications': add_publications_section,
    'languages': add_languages_section,
    'memberships': add_memberships_section,
    'portfolio': add_portfolio_section,
    'published_games': add_published_games_section,
}

def build_resume_document(resume_data):
    """Construct Word document from resume data."""
    doc = Document()

    # Set up style: use narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # =====================================================================
    # TITLE AND CONTACT INFO
    # =====================================================================
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run(resume_data["name"])
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(10, 60, 90)  # Accent blue
    title_run.font.name = 'Calibri'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph(resume_data.get("title", ""))
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if subtitle.runs:
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
        subtitle_run.font.name = 'Calibri'

    # Contact info formatting
    contact = resume_data.get("contact", {})
    if contact:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_items = []
        if "location" in contact:
            contact_items.append(contact["location"])
        if "linkedin" in contact:
            contact_items.append(contact["linkedin"].split("/")[-1])
        if "github" in contact:
            contact_items.append(contact["github"].split("/")[-1])
        contact_para.add_run(" | ".join(contact_items))
        contact_para.paragraph_format.space_before = Pt(6)
        contact_para.paragraph_format.space_after = Pt(12)

    # Build sections in order specified by JSON
    build_sections(doc, resume_data, SECTION_BUILDERS)

    return doc


def main():
    if len(sys.argv) < 2:
        print("Usage: python json2docx.py input.json [output.docx]")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = get_output_path(json_path, sys.argv[2] if len(sys.argv) > 2 else None, '.docx')

    resume_data = load_resume_json(json_path)
    doc = build_resume_document(resume_data)

    if os.path.exists(output_path):
        print(f"Overwriting existing file: {output_path}")

    doc.save(output_path)
    print(f"Wrote Word document to: {output_path}")


if __name__ == "__main__":
    main()
